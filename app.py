import streamlit as st
from google.cloud import translate
import os
import pdfplumber
from docx import Document
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import io
import json

# ---------------- Authentication ---------------- #
if "google_cloud" in st.secrets:
    creds = json.loads(st.secrets["google_cloud"])
    with open("credentials.json", "w") as f:
        json.dump(creds, f)
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = "credentials.json"
else:
    st.error("❌ Google Cloud credentials not found in secrets!")
    st.stop()

# Initialize Google Translate client
translate_client = translate.TranslationServiceClient()

# ---------------- Helper Functions ---------------- #
def extract_text(uploaded_file):
    ext = uploaded_file.name.split('.')[-1].lower()
    text = ""
    if ext == 'txt':
        text = uploaded_file.read().decode('utf-8')
    elif ext == 'pdf':
        with pdfplumber.open(uploaded_file) as pdf:
            for page in pdf.pages:
                page_text = page.extract_text()
                if page_text:
                    text += page_text + "\n"
    elif ext == 'docx':
        doc = Document(uploaded_file)
        for para in doc.paragraphs:
            text += para.text + "\n"
    return text

def translate_text(text, target_lang):
    project_id = st.secrets["project_id"]  # from secrets
    location = "global"

    parent = f"projects/{project_id}/locations/{location}"

    response = translate_client.translate_text(
        request={
            "parent": parent,
            "contents": [text],
            "mime_type": "text/plain",
            "source_language_code": "en",
            "target_language_code": target_lang,
        }
    )

    return response.translations[0].translated_text

def create_docx(translated_text):
    docx_buffer = io.BytesIO()
    doc = Document()
    for line in translated_text.split("\n"):
        doc.add_paragraph(line)
    doc.save(docx_buffer)
    docx_buffer.seek(0)
    return docx_buffer

def create_pdf(translated_text):
    pdf_buffer = io.BytesIO()
    c = canvas.Canvas(pdf_buffer, pagesize=letter)
    width, height = letter
    y = height - 40
    for line in translated_text.split("\n"):
        c.drawString(40, y, line)
        y -= 15
        if y < 40:  # start new page if needed
            c.showPage()
            y = height - 40
    c.save()
    pdf_buffer.seek(0)
    return pdf_buffer

# ---------------- Streamlit UI ---------------- #
st.title("🌍 English → African Languages Translator")

uploaded_file = st.file_uploader("Upload a file (TXT, PDF, DOCX)", type=["txt", "pdf", "docx"])
target_lang = st.text_input("Enter target language code (e.g., 'sw' for Swahili, 'am' for Amharic)")

if uploaded_file and target_lang:
    if st.button("Translate"):
        with st.spinner("Translating..."):
            text = extract_text(uploaded_file)
            if text.strip():
                translated_text = translate_text(text, target_lang)
                st.subheader("✅ Translated Text")
                st.write(translated_text)

                # Download buttons
                docx_file = create_docx(translated_text)
                pdf_file = create_pdf(translated_text)

                st.download_button(
                    label="⬇ Download as DOCX",
                    data=docx_file,
                    file_name="translated_report.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

                st.download_button(
                    label="⬇ Download as PDF",
                    data=pdf_file,
                    file_name="translated_report.pdf",
                    mime="application/pdf"
                )
            else:
                st.warning("No text could be extracted from the file.")
